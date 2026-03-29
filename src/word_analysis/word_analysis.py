"""Module to analyze word documents"""
# import logging
import re
from docx.enum.text import WD_COLOR_INDEX
from ..config import config


class WordAnalyzer:
    """Methods to analyze FQT and SIT data"""
    def acronym_sweep(self, doc, acronyms):
        """Acronym sweep of document"""
        acronyms = acronyms.set_index('Acronym')['Definition'].to_dict()
        acronym_usage = {}
        findings = []
        ignore_words = {'of', 'and', 'the', 'in', 'for', 'on', 'to', 'with', 'by', 'as', 'at', 'from', '&'}
        para_index = 0
        valid_security = config.ConfigOptionsWord["SECURITY_MARKINGS"]
        for section in doc.sections:
            for paragraph in section.header.paragraphs:
                if all(paragraph.text.strip() not in valid_security for paragraph in section.header.paragraphs):
                    findings.append(f"INVALID HEADER SECURITY MARKINGS: {paragraph.text.strip()}")
            for paragraph in section.footer.paragraphs:
                if all(paragraph.text.strip() not in valid_security for paragraph in section.footer.paragraphs):
                    findings.append(f"INVALID FOOTER SECURITY MARKINGS: {paragraph.text.strip()}")
        for para in doc.paragraphs:
            acronym_found_in_para = {}
            if para.style.name.startswith('Heading'):
                # print(f"Header detected: {para.text} (Style: {para.style.name})")
                message = f"________ {para.text} _________"
                # logging.info(message)
                findings.append(message)
                para_index = 0
            for run in para.runs:
                if (run.font.size and run.font.size != 152400) or (run.font.name and run.font.name != 'Arial'):
                    message = f"(Pink) Paragraph {para_index}, position {para.text.find(run.text)}. "\
                              f"{run.text} "
                    run.font.highlight_color = WD_COLOR_INDEX.PINK
                if run.font.size and (run.font.size != '152400'):
                    message += f"(Font Size = {run.font.size / 12700}pt)"
                if run.font.name and (run.font.name != 'Arial'):
                    message += f"(Font = {run.font.name})"
                if (run.font.size and run.font.size != 152400) or (run.font.name and run.font.name != 'Arial'):
                    findings.append(message)
            for acronym, full_form in acronyms.items():
                acronym_pattern = r'\b' + re.escape(acronym.strip()) + r'\b'
                full_form_pattern = r'\b' + re.escape(full_form.strip()) + r'\b'
                for match in re.finditer(acronym_pattern, para.text):
                    position = match.start()
                    if acronym not in acronym_usage:
                        if not re.search(full_form_pattern, para.text):
                            message = f"(Yellow) Paragraph {para_index}, position {position}. "\
                                      f"First instance of {acronym} should have the definition '{full_form}' before it."
                            # logging.warning(message)
                            findings.append(message)
                            self.highlight_text(para, position, len(acronym), WD_COLOR_INDEX.DARK_YELLOW)
                            # para_parts = para.text.split(acronym)
                            # para.text = para_parts[0] + f" {full_form} ({acronym})" + ''.join(para_parts[1:])
                        acronym_usage[acronym] = 'first'
                        acronym_found_in_para[acronym] = True
                    else:
                        if re.search(full_form_pattern, para.text):
                            message = f"(Teal) Paragraph {para_index}, position {para.text.find(full_form)}. "\
                                      f"{full_form} ({acronym}) has already been defined, replace with the acronym."
                            # logging.warning(message)
                            findings.append(message)
                            self.highlight_text(para, para.text.find(full_form), len(full_form), WD_COLOR_INDEX.TEAL)
                            # para.text = re.sub(acronym_pattern, '', para.text)
                            # para.text = para.text.replace('()', '')
                            # para.text = re.sub(full_form_pattern, acronym, para.text)
                        acronym_usage[acronym] = 'used'
                for match in re.finditer(full_form_pattern, para.text):
                    position = match.start()
                    if acronym in acronym_usage and acronym not in acronym_found_in_para:
                        message = f"(Green) Paragraph {para_index}, position {position}. "\
                                  f"{full_form} should be replaced with {acronym} in subsequent uses."
                        # logging.warning(message)
                        findings.append(message)
                        self.highlight_text(para, position, len(full_form), WD_COLOR_INDEX.BRIGHT_GREEN)
                        # para.text = re.sub(full_form_pattern, acronym, para.text)
            potential_acronym = []
            capitalized_word_count = 0
            for position, word in enumerate(para.text.split()):
                cleaned_word = re.sub(r'[^\w/&]', '', word)
                if cleaned_word.isupper() and cleaned_word not in acronyms and len(cleaned_word) > 1:
                    message = f"(Violet) Paragraph {para_index}, position {para.text.find(word)}. "\
                              f"Found acronym {cleaned_word} not in acronym list."
                    # logging.warning(message)
                    findings.append(message)
                    self.highlight_text(para, para.text.find(word), len(word), WD_COLOR_INDEX.VIOLET)
                if word[0].isupper() and not word.isupper():
                    potential_acronym.append(word)
                    capitalized_word_count += 1
                elif len(potential_acronym) > 0 and word.lower() in ignore_words:
                    potential_acronym.append(word)
                elif len(potential_acronym) > 1 and capitalized_word_count > 1:
                    acronym_candidate = ' '.join(potential_acronym)
                    if acronym_candidate not in acronyms.values():
                        position = para.text.find(' '.join(potential_acronym))
                        message = f"(Blue) Paragraph {para_index}, position {position}. Found potential acronym: '{acronym_candidate}'."
                        # logging.warning(message)
                        findings.append(message)
                        self.highlight_text(para, position, len(acronym_candidate), WD_COLOR_INDEX.BLUE)
                    potential_acronym = []
                    capitalized_word_count = 0
                else:
                    potential_acronym = []
                    capitalized_word_count = 0
            if len(potential_acronym) > 1:
                acronym_candidate = ' '.join(potential_acronym)
                if acronym_candidate not in acronyms.values():
                    position = para.text.find(' '.join(potential_acronym))
                    message = f"(Blue) Paragraph {para_index}, position {position}. Found potential acronym: '{acronym_candidate}'."
                    # logging.warning(message)
                    findings.append(message)
                    self.highlight_text(para, position, len(acronym_candidate), WD_COLOR_INDEX.BLUE)
            para_index += 1
            index = 0
            while index < len(para.text) - 1:
                if para.text[index:index + 2] == "  ":
                    if index == 0 or para.text[index - 1] != ".":
                        message = f"(Red) Double space found at position {index} in paragraph {para_index}: {para.text[index-7:index+7]}"
                        findings.append(message)
                        self.highlight_text(para, index, 2, WD_COLOR_INDEX.RED)
                        # logging.warning(message)
                index += 1
        for acronym in acronyms.keys():
            if acronym not in acronym_usage:
                message = f"Acronym {acronym} in table but not used"
                # logging.warning(message)
                findings.append(message)
        for acronym, usage in acronym_usage.items():
            if usage == 'first':
                message = f"Acronym {acronym} defined but not used"
                # logging.warning(message)
                findings.append(message)
        for number, table in enumerate(doc.tables):
            findings.append(f"------ Table {number} ------")
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        for run in para.runs:
                            if (run.font.size and run.font.size != 127000) or (run.font.name and run.font.name != 'Arial'):
                                message = f"(Pink) Paragraph {para_index}, position {para.text.find(run.text)}. "\
                                          f"{run.text} "
                                run.font.highlight_color = WD_COLOR_INDEX.PINK
                            if run.font.size and (run.font.size != 127000):
                                message += f"(Font Size = {run.font.size / 12700}pt)"
                            if run.font.name and (run.font.name != 'Arial'):
                                message += f"(Font = {run.font.name})"
                            if (run.font.size and run.font.size != 127000) or (run.font.name and run.font.name != 'Arial'):
                                findings.append(message)
        return doc, findings

    def highlight_text(self, para, position, length, color):
        """Find the correct run, split it if needed, and highlight the issue"""
        current_pos = 0
        # run_count = len(para.runs)
        # print(run_count)
        for run in para.runs:
            run_length = len(run.text)
            font = run.font.name
            size = run.font.size
            highlight = run.font.highlight_color
            delete_element = run._element
            if current_pos <= position < current_pos + run_length:
                offset = position - current_pos
                affected_text = run.text[offset:offset + length]
                before_text = run.text[:offset]
                after_text = run.text[offset + length:]
                if before_text:
                    new_run = para.add_run(before_text)
                    new_run.font.name = font
                    new_run.font.size = size
                    new_run.font.highlight_color = highlight

                if affected_text:
                    highlighted_run = para.add_run(affected_text)
                    highlighted_run.font.name = font
                    highlighted_run.font.size = size
                    highlighted_run.font.highlight_color = color

                if after_text:
                    new_run = para.add_run(after_text)
                    new_run.font.name = font
                    new_run.font.size = size
                    new_run.font.highlight_color = highlight
            else:
                new_run = para.add_run(run.text)
                new_run.font.name = font
                new_run.font.size = size
                new_run.font.highlight_color = highlight
            current_pos += run_length
            para._element.remove(delete_element)

    def placeholder(self):
        """Placeholder"""
        return
